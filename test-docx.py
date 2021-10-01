# Import lib MSWord and Wikipedia
from docx import Document
import wikipedia


def wiki(keyword, lang = "th"):
    wikipedia.set_lang(lang)

    # summary เอาบทความสรุป
    data = wikipedia.summary(keyword)

    # page + content เอาบทความทั้้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    # สร้าง และจัดการDocument
    doc = Document()
    
    doc.add_heading(keyword,0)
    doc.add_paragraph(data2)

    doc.save(keyword+".docx")

    print("Completed generating",keyword,"file")


try:
    wiki("ewqecwqeqwe","en")
except:
    print("Not found, please try again")
# wiki("ญี่ปุ่น","en")














































