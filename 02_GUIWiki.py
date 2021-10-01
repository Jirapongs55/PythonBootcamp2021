# Wiki
import wikipedia

# python to docx    
from docx import Document


def wiki(keyword, lang = "th"):
    wikipedia.set_lang(lang)

    # summary เอาแค่สรุปบทความ
    data = wikipedia.summary(keyword)
    # page + content เอาเนื้อหาทั้งหน้าบทความ
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    # ส่งข้อมูล และจัดการในไฟล์ docx
    docx = Document()   # สร้างไฟล์ word
    docx.add_heading(keyword,0)
    docx.add_paragraph(data2)
    docx.save(keyword + ".docx")
    print("Completed generating",keyword,"file")        
    

# เปลี่ยนภาษา
wikipedia.set_lang("th")

# GUI
from tkinter import *
from tkinter import ttk
from tkinter import messagebox


GUI = Tk()
GUI.title("โปรแกรม wikipedia by K'Jirapong")
GUI.geometry("600x300")


# Config
font_1 = ("TH Sarabun New",16)


# คำอธิบาย
L1 = ttk.Label(GUI, text = "กรุณากรอกสิ่งที่ต้องการค้นหา", font = font_1)
L1.pack()


# ช่องค้นหาข้อมูล
v_search = StringVar()  # กล่องสำหรับเก็บ keyword
E1 = ttk.Entry(GUI, textvariable = v_search, width = 42, font = font_1)
E1.pack(pady = 10)


def search():
    keyword = v_search.get()    # .get() คือคำสั่งในการดึงข้อมูลจากตัวแปรเข้ามา (ใช้ได้เฉพาะ StringVar()เท่านั้น !!! IntVar(), FloatVar() ใช้ไม่ได้)
    try:    # ลองค้นหาดูว่าได้ผลลัพธ์หรือไม่ หากได้ให้ผ่านไป else
        lang = v_radio.get()    # th/en/zh
        wiki(keyword, lang)
    except: # หากไม่ได้หรือมีปัญหา ให้แสดง message box แจ้งเตือนuser
        messagebox.showwarning("Alert message", "please try keyword again")
    else:   # หากได้ต้องการให้ขึ้น message box แสดงว่าเสร็จแล้ว
        messagebox.showinfo("Completed message", "Completed generating the " + keyword + " file")
    '''
    print(wikipedia.search(keyword))
    result = wikipedia.summary(keyword)
    print(result)
    '''
    

# ปุ่มค้นหา
B1 = ttk.Button(GUI, text = "Search", command = search)
B1.pack(ipadx = 20, ipady = 10, pady =10)


# เมนูเลือกภาษา
F1 = Frame(GUI)
F1.pack()


v_radio = StringVar()   # เป็นช่องเก็บข้อมูลภาษา
RB1 = ttk.Radiobutton(F1, text = "ภาษาไทย", variable = v_radio, value = "th")
RB2 = ttk.Radiobutton(F1, text = "ภาษาอังกฤษ", variable = v_radio, value = "en")
RB3 = ttk.Radiobutton(F1, text = "ภาษาจีน", variable = v_radio, value = "zh")
RB1.invoke()

RB1.grid(row = 0, column = 0)
RB2.grid(row = 0, column = 1)
RB3.grid(row = 0, column = 2)


GUI.mainloop()











































































